from __future__ import annotations

import re
from pathlib import Path
from urllib.parse import urlparse
from xml.sax.saxutils import escape

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown_it import MarkdownIt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import HRFlowable, LongTable, PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer, TableStyle

ROOT = Path(__file__).resolve().parents[1]
CHAPTERS_DIR = ROOT / "chapters"
FRONT_MATTER_DIR = ROOT / "front_matter"
REFERENCE_DIR = ROOT / "reference"
BUILD_DIR = ROOT / "build"
CONFIG_PATH = ROOT / "book.yml"
FINAL_CHAPTER = 30
REFERENCE_BEFORE_CHAPTER = 29
FINAL_LINE = "**L'AI può scrivere il codice. Il timone resta a noi.**"
SECTION_HEADING_RE = re.compile(r"^\d+\.\d+(?:\s|\b)")
URL_RE = re.compile(r"https?://[^\s)>]+")
REAL_CASE_RE = re.compile(r"\b(?:il\s+|un\s+)?caso reale(?: documentato)?\b", re.IGNORECASE)
FOOTNOTE_DEF_RE = re.compile(r"(?m)^\[\^([^\]]+)\]:\s*(.+?)\s*$")
FOOTNOTE_REF_RE = re.compile(r"\[\^([^\]]+)\]")
READER_REFERENCE_RE = re.compile(r"^\d+.*\.md$", re.IGNORECASE)
ACCENT_HEX = "176B68"
ACCENT = colors.HexColor(f"#{ACCENT_HEX}")
INK = colors.HexColor("#172326")
MUTED = colors.HexColor("#526065")
SOFT = colors.HexColor("#EEF4F3")
ACCENT_RGB = RGBColor(0x17, 0x6B, 0x68)
INK_RGB = RGBColor(0x17, 0x23, 0x26)


def load_config() -> dict:
    with CONFIG_PATH.open("r", encoding="utf-8") as handle:
        config = yaml.safe_load(handle) or {}
    required = ("title", "subtitle", "author", "output_basename", "language")
    missing = [key for key in required if not str(config.get(key, "")).strip()]
    if missing:
        raise ValueError(f"book.yml: metadata mancanti: {', '.join(missing)}")
    return config


def numeric_prefix(path: Path) -> int:
    match = re.match(r"(\d+)", path.name)
    return int(match.group(1)) if match else 10**9


def source_sort_key(path: Path) -> tuple[int, str]:
    return numeric_prefix(path), path.name.casefold()


def front_matter_files() -> list[Path]:
    if not FRONT_MATTER_DIR.exists():
        return []
    return sorted(FRONT_MATTER_DIR.glob("*.md"), key=source_sort_key)


def reference_files() -> list[Path]:
    if not REFERENCE_DIR.exists():
        return []
    return sorted([p for p in REFERENCE_DIR.glob("*.md") if READER_REFERENCE_RE.match(p.name)], key=source_sort_key)


def chapter_dirs() -> list[Path]:
    return sorted([p for p in CHAPTERS_DIR.glob("*_chapter") if p.is_dir()], key=source_sort_key)


def chapter_files(chapter_dir: Path) -> list[Path]:
    return sorted(chapter_dir.glob("*.md"), key=source_sort_key)


def source_files() -> list[Path]:
    return [path for directory in chapter_dirs() for path in chapter_files(directory)]


def resolve_footnotes(text: str, scope: str) -> str:
    definitions: dict[str, str] = {}
    for match in FOOTNOTE_DEF_RE.finditer(text):
        key = match.group(1)
        value = match.group(2).strip()
        current = definitions.get(key)
        if current is None or len(value) > len(current):
            definitions[key] = value
    body = FOOTNOTE_DEF_RE.sub("", text)
    numbering: dict[str, int] = {}
    missing: set[str] = set()

    def repl(match: re.Match[str]) -> str:
        key = match.group(1)
        if key not in definitions:
            missing.add(key)
            return match.group(0)
        numbering.setdefault(key, len(numbering) + 1)
        return f"[{numbering[key]}]"

    body = FOOTNOTE_REF_RE.sub(repl, body)
    if missing:
        raise ValueError(f"{scope}: note senza definizione: {', '.join(sorted(missing))}")
    body = re.sub(r"\n{3,}", "\n\n", body).strip()
    if not numbering:
        return body + "\n\n"
    notes = ["## Note e fonti", ""]
    for key, number in sorted(numbering.items(), key=lambda item: item[1]):
        notes.append(f"{number}. {definitions[key]}")
    return body + "\n\n" + "\n".join(notes).strip() + "\n\n"


def first_h1(path: Path) -> str:
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.startswith("# "):
            return line[2:].strip()
    return path.stem


def chapter_titles(files: list[Path]) -> dict[Path, str]:
    titles: dict[Path, str] = {}
    for path in files:
        titles.setdefault(path.parent, first_h1(path))
    return titles


def chapter_index(files: list[Path]) -> str:
    entries = [f"- {title}" for title in chapter_titles(files).values()]
    return "# Indice dei capitoli\n\n" + "\n".join(entries) + "\n\n"


def real_case_index(files: list[Path]) -> str:
    titles = chapter_titles(files)
    grouped: dict[str, list[str]] = {}
    seen: set[tuple[str, str]] = set()
    for path in files:
        chapter = titles[path.parent]
        for line in path.read_text(encoding="utf-8").splitlines():
            stripped = line.strip()
            if not re.match(r"^#{1,6}\s+", stripped):
                continue
            heading = re.sub(r"^#{1,6}\s+", "", stripped).strip()
            case_match = REAL_CASE_RE.search(heading)
            if not case_match:
                continue
            label = re.sub(r"^[\s—:–-]+", "", heading[case_match.end():]).strip()
            if not label:
                continue
            key = (chapter, label.casefold())
            if key in seen:
                continue
            seen.add(key)
            grouped.setdefault(chapter, []).append(label)
    chunks = ["# Indice dei casi reali documentati\n\n", "Questo indice raccoglie i casi che il manoscritto identifica esplicitamente come reali. Gli scenari Example Software Industries (ESI) sono fittizi/compositi e non sono inclusi.\n\n"]
    for chapter, labels in grouped.items():
        chunks.append(f"## {chapter}\n\n")
        chunks.extend(f"- {label}\n" for label in labels)
        chunks.append("\n")
    return "".join(chunks)


def source_label(line: str, url: str) -> str:
    label = line.replace(url, " ")
    label = re.sub(r"^\s*\[\^[^\]]+\]:\s*", "", label).strip().lstrip("- ").strip().rstrip(" :;,.—-")
    if URL_RE.search(label) or label.casefold() in {"fonte", "fonti", "url", "url canonico"}:
        return ""
    return (label[:177].rstrip() + "...") if len(label) > 180 else label


def source_index(files: list[Path]) -> str:
    titles = chapter_titles(files)
    seen_urls: set[str] = set()
    entries: list[tuple[str, str, str, str, str]] = []
    for path in files:
        chapter = titles[path.parent]
        for line in path.read_text(encoding="utf-8").splitlines():
            for match in URL_RE.findall(line):
                url = match.rstrip(".,;:")
                if url in seen_urls:
                    continue
                seen_urls.add(url)
                domain = urlparse(url).netloc.removeprefix("www.")
                label = source_label(line, url) or domain
                entries.append((domain.casefold(), label.casefold(), label, chapter, url))
    entries.sort(key=lambda item: (item[0], item[1], item[4]))
    chunks = ["# Indice delle fonti\n\n", f"Questo indice consolida **{len(entries)} URL esterni distinti** presenti nel corpo del libro. Le fonti restano vicine ai claim che sostengono; qui sono raccolte per facilitarne il ritrovamento.\n\n"]
    current_domain: str | None = None
    for _, _, label, chapter, url in entries:
        domain = urlparse(url).netloc.removeprefix("www.")
        if domain != current_domain:
            current_domain = domain
            chunks.append(f"## {domain}\n\n")
        chunks.append(f"- {label} — [{domain}]({url}) — {chapter}\n")
    chunks.append("\n")
    return "".join(chunks)


def assemble_chapter(directory: Path) -> str:
    raw = "\n\n".join(path.read_text(encoding="utf-8").strip() for path in chapter_files(directory))
    return resolve_footnotes(raw, directory.name)


def assemble_markdown(config: dict) -> str:
    all_files = source_files()
    by_number = {numeric_prefix(directory): directory for directory in chapter_dirs()}
    missing = sorted(set(range(FINAL_CHAPTER + 1)) - set(by_number))
    if missing:
        raise ValueError(f"Capitoli mancanti: {missing}")
    chunks = [f"# {config['title']}\n\n## {config['subtitle']}\n\n**Autore:** {config['author']}\n\n---\n\n"]
    for path in front_matter_files():
        chunks.append(resolve_footnotes(path.read_text(encoding="utf-8"), path.name))
    chunks.append(chapter_index(all_files))
    for number in range(REFERENCE_BEFORE_CHAPTER):
        chunks.append(assemble_chapter(by_number[number]))
    for path in reference_files():
        chunks.append(resolve_footnotes(path.read_text(encoding="utf-8"), path.name))
    chunks.append(real_case_index(all_files))
    chunks.append(source_index(all_files))
    chunks.append(assemble_chapter(by_number[29]))
    chunks.append(assemble_chapter(by_number[30]).rstrip())
    assembled = "".join(chunks).rstrip() + "\n"
    last_nonempty = next(line for line in reversed(assembled.splitlines()) if line.strip())
    if last_nonempty != FINAL_LINE:
        raise ValueError(f"Invariante finale violata: trovata {last_nonempty!r}")
    return assembled


def markdown_parser() -> MarkdownIt:
    return MarkdownIt("commonmark").enable("table")


def effective_heading_level(level: int, content: str) -> int:
    return 2 if level == 1 and SECTION_HEADING_RE.match(content.strip()) else level


def inline_plain(token) -> str:
    if not getattr(token, "children", None):
        return token.content
    parts: list[str] = []
    for child in token.children:
        if child.type in {"text", "code_inline", "html_inline"}:
            parts.append(child.content)
        elif child.type in {"softbreak", "hardbreak"}:
            parts.append("\n")
        elif child.type == "image":
            parts.append(child.content or child.attrGet("alt") or "")
    return "".join(parts)


def add_docx_hyperlink(paragraph, text: str, url: str) -> None:
    r_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT_HEX)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.extend([r_pr, text_element])
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_docx_inline(paragraph, token) -> None:
    children = getattr(token, "children", None)
    if not children:
        paragraph.add_run(token.content)
        return
    bold_depth = italic_depth = 0
    link_href: str | None = None
    for child in children:
        if child.type == "strong_open":
            bold_depth += 1; continue
        if child.type == "strong_close":
            bold_depth = max(0, bold_depth - 1); continue
        if child.type == "em_open":
            italic_depth += 1; continue
        if child.type == "em_close":
            italic_depth = max(0, italic_depth - 1); continue
        if child.type == "link_open":
            link_href = child.attrGet("href") or None; continue
        if child.type == "link_close":
            link_href = None; continue
        if child.type in {"softbreak", "hardbreak"}:
            paragraph.add_run("\n"); continue
        if child.type not in {"text", "code_inline", "html_inline", "image"}:
            continue
        text = child.content if child.type != "image" else (child.content or child.attrGet("alt") or "")
        if link_href and child.type == "text":
            add_docx_hyperlink(paragraph, text, link_href); continue
        run = paragraph.add_run(text)
        run.bold = bold_depth > 0
        run.italic = italic_depth > 0
        if child.type == "code_inline":
            run.font.name = "Consolas"; run.font.size = Pt(9.2)


def inline_reportlab(token) -> str:
    children = getattr(token, "children", None)
    if not children:
        return escape(token.content)
    parts: list[str] = []
    for child in children:
        if child.type in {"text", "html_inline"}: parts.append(escape(child.content))
        elif child.type == "code_inline": parts.append(f'<font name="Courier">{escape(child.content)}</font>')
        elif child.type == "strong_open": parts.append("<b>")
        elif child.type == "strong_close": parts.append("</b>")
        elif child.type == "em_open": parts.append("<i>")
        elif child.type == "em_close": parts.append("</i>")
        elif child.type in {"softbreak", "hardbreak"}: parts.append("<br/>")
        elif child.type == "link_open": parts.append(f'<link href="{escape(child.attrGet("href") or "")}" color="#{ACCENT_HEX}">')
        elif child.type == "link_close": parts.append("</link>")
        elif child.type == "image": parts.append(escape(child.content or child.attrGet("alt") or ""))
    return "".join(parts)


def parse_table(tokens, start: int, render_cell) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    row: list[str] | None = None
    i = start + 1
    while i < len(tokens):
        token = tokens[i]
        if token.type == "table_close": return rows, i + 1
        if token.type == "tr_open": row = []
        elif token.type == "tr_close":
            if row is not None: rows.append(row)
            row = None
        elif token.type in {"th_open", "td_open"} and row is not None and i + 1 < len(tokens) and tokens[i + 1].type == "inline":
            row.append(render_cell(tokens[i + 1]))
        i += 1
    return rows, i


def mark_docx_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader"); header.set(qn("w:val"), "true"); tr_pr.append(header)


def add_docx_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_docx_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"; normal.font.size = Pt(10.8); normal.font.color.rgb = INK_RGB
    normal.paragraph_format.space_after = Pt(6.5); normal.paragraph_format.line_spacing = 1.17
    title = doc.styles["Title"]
    title.font.name = "Aptos Display"; title.font.size = Pt(34); title.font.bold = True; title.font.color.rgb = ACCENT_RGB; title.paragraph_format.space_after = Pt(18)
    for name, size in (("Heading 1", 28), ("Heading 2", 18), ("Heading 3", 14), ("Heading 4", 11.5)):
        style = doc.styles[name]
        style.font.name = "Aptos Display"; style.font.size = Pt(size); style.font.bold = True
        style.font.color.rgb = ACCENT_RGB if name != "Heading 4" else INK_RGB
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(18 if name == "Heading 1" else 12); style.paragraph_format.space_after = Pt(9 if name != "Heading 4" else 5)
    code = doc.styles.add_style("Book Code", 1) if "Book Code" not in doc.styles else doc.styles["Book Code"]
    code.font.name = "Consolas"; code.font.size = Pt(8.8); code.paragraph_format.left_indent = Inches(0.2); code.paragraph_format.right_indent = Inches(0.1); code.paragraph_format.space_before = Pt(5); code.paragraph_format.space_after = Pt(8)
    section = doc.sections[0]
    section.top_margin = Inches(0.75); section.bottom_margin = Inches(0.72); section.left_margin = Inches(0.82); section.right_margin = Inches(0.82); section.different_first_page_header_footer = True


def build_docx(markdown: str, output: Path, config: dict) -> None:
    tokens = markdown_parser().parse(markdown)
    doc = Document(); configure_docx_styles(doc)
    doc.core_properties.title = config["title"]; doc.core_properties.subject = config["subtitle"]; doc.core_properties.author = config["author"]; doc.core_properties.last_modified_by = config["author"]
    i = 0; first_h1 = True; title_page = True; list_stack: list[str] = []; blockquote_depth = 0
    while i < len(tokens):
        token = tokens[i]
        if token.type == "heading_open":
            content_token = tokens[i + 1]; content = inline_plain(content_token); level = effective_heading_level(int(token.tag[1]), content)
            if first_h1:
                p = doc.add_paragraph(style="Title"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(100); add_docx_inline(p, content_token); first_h1 = False
            else:
                if level == 1: doc.add_page_break(); title_page = False
                p = doc.add_heading(level=min(level, 4)); add_docx_inline(p, content_token)
            i += 3; continue
        if token.type == "paragraph_open":
            content_token = tokens[i + 1]; plain = inline_plain(content_token).strip()
            style = "Quote" if blockquote_depth else ("List Bullet" if list_stack and list_stack[-1] == "bullet" else "List Number" if list_stack else None)
            p = doc.add_paragraph(style=style)
            if title_page and plain.startswith("Autore:"): p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(18)
            add_docx_inline(p, content_token); i += 3; continue
        if token.type == "table_open":
            rows, i = parse_table(tokens, i, inline_plain)
            if rows:
                cols = max(len(row) for row in rows); table = doc.add_table(rows=len(rows), cols=cols); table.style = "Table Grid"; table.autofit = True
                for r_idx, row in enumerate(rows):
                    for c_idx, value in enumerate(row):
                        cell = table.cell(r_idx, c_idx); cell.text = value
                        if r_idx == 0:
                            for run in cell.paragraphs[0].runs: run.bold = True
                mark_docx_header_row(table.rows[0]); doc.add_paragraph()
            continue
        if token.type == "blockquote_open": blockquote_depth += 1
        elif token.type == "blockquote_close": blockquote_depth = max(0, blockquote_depth - 1)
        elif token.type == "bullet_list_open": list_stack.append("bullet")
        elif token.type == "ordered_list_open": list_stack.append("number")
        elif token.type in {"bullet_list_close", "ordered_list_close"} and list_stack: list_stack.pop()
        elif token.type == "fence":
            p = doc.add_paragraph(style="Book Code"); run = p.add_run(token.content.rstrip()); run.font.name = "Consolas"; run.font.size = Pt(8.8)
        elif token.type == "hr":
            p = doc.add_paragraph("·  ·  ·"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
    section = doc.sections[0]; section.different_first_page_header_footer = True; section.first_page_footer.paragraphs[0].text = ""; add_docx_page_field(section.footer.paragraphs[0]); doc.save(output)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="BookTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=31, leading=36, alignment=TA_CENTER, textColor=ACCENT, spaceBefore=95, spaceAfter=20))
    styles.add(ParagraphStyle(name="BookSubtitle", parent=styles["BodyText"], fontName="Helvetica-Oblique", fontSize=14, leading=19, alignment=TA_CENTER, textColor=ACCENT, spaceAfter=14))
    styles.add(ParagraphStyle(name="BookAuthor", parent=styles["BodyText"], fontName="Helvetica", fontSize=11.5, leading=16, alignment=TA_CENTER, textColor=INK, spaceAfter=12))
    styles.add(ParagraphStyle(name="H1Book", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=24, leading=29, textColor=ACCENT, spaceBefore=18, spaceAfter=15, keepWithNext=True))
    styles.add(ParagraphStyle(name="H2Book", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=15.5, leading=19.5, textColor=ACCENT, spaceBefore=13, spaceAfter=8, keepWithNext=True))
    styles.add(ParagraphStyle(name="H3Book", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=12.2, leading=15.5, textColor=INK, spaceBefore=10, spaceAfter=6, keepWithNext=True))
    styles.add(ParagraphStyle(name="BodyBook", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.25, leading=14.8, textColor=INK, spaceAfter=7))
    styles.add(ParagraphStyle(name="QuoteBook", parent=styles["BodyText"], fontName="Helvetica-Oblique", fontSize=10.25, leading=14.8, textColor=INK, leftIndent=10*mm, rightIndent=6*mm, borderColor=ACCENT, borderWidth=1.5, borderPadding=(2,0,2,7), spaceAfter=8))
    styles.add(ParagraphStyle(name="TableBook", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.7, leading=9.4, textColor=INK, spaceAfter=0, wordWrap="CJK"))
    styles.add(ParagraphStyle(name="CodeBook", parent=styles["Code"], fontName="Courier", fontSize=7.8, leading=10.1, leftIndent=5*mm, rightIndent=3*mm, backColor=SOFT, borderColor=ACCENT, borderWidth=0.6, borderPadding=5, spaceBefore=4, spaceAfter=7))
    return styles


class BookPdfTemplate(SimpleDocTemplate):
    def afterFlowable(self, flowable) -> None:
        key = getattr(flowable, "_bookmark_key", None)
        if not key: return
        self.canv.bookmarkPage(key); self.canv.addOutlineEntry(getattr(flowable, "_bookmark_title", ""), key, level=getattr(flowable, "_bookmark_level", 0), closed=False)


def build_pdf(markdown: str, output: Path, config: dict) -> None:
    tokens = markdown_parser().parse(markdown); styles = pdf_styles(); story = []
    i = 0; first_h1 = True; title_page = True; list_stack: list[dict[str, int | str]] = []; blockquote_depth = 0; bookmark_counter = 0
    def make_heading(content: str, plain: str, style, level: int) -> Paragraph:
        nonlocal bookmark_counter
        p = Paragraph(content, style)
        if level == 1 or (level == 2 and SECTION_HEADING_RE.match(plain.strip()) is not None):
            bookmark_counter += 1; p._bookmark_key = f"bookmark-{bookmark_counter}"; p._bookmark_title = plain; p._bookmark_level = 0 if level == 1 else 1
        return p
    while i < len(tokens):
        token = tokens[i]
        if token.type == "heading_open":
            content_token = tokens[i + 1]; plain = inline_plain(content_token); content = inline_reportlab(content_token); level = effective_heading_level(int(token.tag[1]), plain)
            if first_h1: story.append(make_heading(content, plain, styles["BookTitle"], 1)); first_h1 = False
            elif level == 1: story.append(PageBreak()); title_page = False; story.append(make_heading(content, plain, styles["H1Book"], 1))
            elif title_page and level == 2: story.append(Paragraph(content, styles["BookSubtitle"]))
            elif level == 2: story.append(make_heading(content, plain, styles["H2Book"], 2))
            else: story.append(Paragraph(content, styles["H3Book"]))
            i += 3; continue
        if token.type == "paragraph_open":
            content_token = tokens[i + 1]; plain = inline_plain(content_token).strip(); content = inline_reportlab(content_token)
            if title_page and plain.startswith("Autore:"): story.append(Paragraph(content, styles["BookAuthor"])); i += 3; continue
            if list_stack:
                current = list_stack[-1]
                if current["type"] == "bullet": content = "&#8226; " + content
                else: current["counter"] = int(current["counter"]) + 1; content = f"{current['counter']}. " + content
            story.append(Paragraph(content, styles["QuoteBook"] if blockquote_depth else styles["BodyBook"])); i += 3; continue
        if token.type == "table_open":
            rows, i = parse_table(tokens, i, inline_reportlab)
            if rows:
                cols = max(len(row) for row in rows); normalized = [row + [""]*(cols-len(row)) for row in rows]; data = [[Paragraph(cell, styles["TableBook"]) for cell in row] for row in normalized]
                usable_width = A4[0] - 36*mm; table = LongTable(data, colWidths=[usable_width/cols]*cols, repeatRows=1, hAlign="LEFT")
                table.setStyle(TableStyle([("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#A8B7B5")),("BACKGROUND",(0,0),(-1,0),SOFT),("TEXTCOLOR",(0,0),(-1,0),INK),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),3.5),("RIGHTPADDING",(0,0),(-1,-1),3.5),("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3)])); story.append(table); story.append(Spacer(1,7))
            continue
        if token.type == "blockquote_open": blockquote_depth += 1
        elif token.type == "blockquote_close": blockquote_depth = max(0, blockquote_depth-1)
        elif token.type == "bullet_list_open": list_stack.append({"type":"bullet","counter":0})
        elif token.type == "ordered_list_open": list_stack.append({"type":"number","counter":0})
        elif token.type in {"bullet_list_close","ordered_list_close"} and list_stack: list_stack.pop()
        elif token.type == "fence": story.append(Preformatted(token.content.rstrip(), styles["CodeBook"], maxLineLength=92, splitChars=" /._-:;,+=")); story.append(Spacer(1,4))
        elif token.type == "hr": story.append(Spacer(1,5)); story.append(HRFlowable(width="30%", thickness=0.7, color=colors.HexColor("#9DB7B4"), hAlign="CENTER")); story.append(Spacer(1,7))
        i += 1
    def set_metadata(canvas): canvas.setTitle(config["title"]); canvas.setAuthor(config["author"]); canvas.setSubject(config["subtitle"])
    def first_page(canvas, doc): set_metadata(canvas)
    def later_pages(canvas, doc):
        set_metadata(canvas); canvas.saveState(); canvas.setFont("Helvetica",8); canvas.setFillColor(MUTED); canvas.drawCentredString(A4[0]/2,10.5*mm,str(doc.page)); canvas.restoreState()
    pdf = BookPdfTemplate(str(output), pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=18*mm, bottomMargin=18*mm, title=config["title"], author=config["author"], subject=config["subtitle"])
    pdf.build(story, onFirstPage=first_page, onLaterPages=later_pages)


def main() -> None:
    config = load_config()
    if not source_files(): raise SystemExit("Nessun file Markdown trovato in chapters/.")
    BUILD_DIR.mkdir(parents=True, exist_ok=True); basename = config["output_basename"]; markdown = assemble_markdown(config)
    md_path = BUILD_DIR / f"{basename}.md"; docx_path = BUILD_DIR / f"{basename}.docx"; pdf_path = BUILD_DIR / f"{basename}.pdf"
    md_path.write_text(markdown, encoding="utf-8"); build_docx(markdown, docx_path, config); build_pdf(markdown, pdf_path, config)
    print(f"Build completata: {len(chapter_dirs())} capitoli, {len(source_files())} file Markdown corpo.")
    print(f"- {md_path.relative_to(ROOT)}\n- {docx_path.relative_to(ROOT)}\n- {pdf_path.relative_to(ROOT)}")


if __name__ == "__main__": main()
