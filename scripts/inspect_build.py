from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import yaml
from docx import Document
from docx.oxml.ns import qn
from markdown_it import MarkdownIt
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
BUILD_DIR = ROOT / "build"
CONFIG = yaml.safe_load((ROOT / "book.yml").read_text(encoding="utf-8")) or {}
BASENAME = CONFIG["output_basename"]
FINAL_TEXT = "L'AI può scrivere il codice. Il timone resta a noi."
FINAL_MARKDOWN_LINE = f"**{FINAL_TEXT}**"
CHAPTER_30 = "Capitolo 30 — I Dieci comandamenti della Software Architecture nell'era dell'AI"
CHAPTER_HEADING_RE = re.compile(r"^Capitolo (\d+)\s+—\s+")


def flatten_outline(items):
    flattened = []
    for item in items or []:
        if isinstance(item, list):
            flattened.extend(flatten_outline(item))
        else:
            flattened.append(str(getattr(item, "title", item)))
    return flattened


def markdown_chapter_numbers(text: str) -> list[str]:
    tokens = MarkdownIt("commonmark").enable("table").parse(text)
    numbers: list[str] = []
    for index, token in enumerate(tokens):
        if token.type != "heading_open" or token.tag != "h1" or index + 1 >= len(tokens):
            continue
        inline = tokens[index + 1]
        if inline.type != "inline":
            continue
        match = CHAPTER_HEADING_RE.match(inline.content.strip())
        if match:
            numbers.append(match.group(1))
    return numbers


def inspect_markdown(path: Path) -> dict:
    text = path.read_text(encoding="utf-8")
    chapter_headings = markdown_chapter_numbers(text)
    last = next(line for line in reversed(text.splitlines()) if line.strip())
    assert CONFIG["title"] in text and CONFIG["author"] in text
    assert chapter_headings == [str(i) for i in range(31)], chapter_headings
    assert "[^" not in text, "Footnote Markdown irrisolte"
    assert last == FINAL_MARKDOWN_LINE, last
    assert text.rfind(CHAPTER_30) > text.rfind("# Indice delle fonti")
    return {"chapters": len(chapter_headings), "bytes": path.stat().st_size}


def inspect_docx(path: Path) -> dict:
    doc = Document(path)
    assert doc.core_properties.title == CONFIG["title"]
    assert doc.core_properties.subject == CONFIG["subtitle"]
    assert doc.core_properties.author == CONFIG["author"]
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert any(CHAPTER_30 in p for p in paragraphs)
    assert paragraphs[-1] == FINAL_TEXT, paragraphs[-1]
    assert not any(re.search(r"\[\^[^\]]+\]", p) for p in paragraphs)
    assert doc.sections[0].first_page_footer.paragraphs[0].text.strip() == ""
    tables = len(doc.tables)
    repeated = 0
    for table in doc.tables:
        if table.rows and table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None:
            repeated += 1
    assert repeated == tables, f"Repeating header {repeated}/{tables}"
    return {"tables": tables, "repeating_headers": repeated, "bytes": path.stat().st_size}


def inspect_pdf(path: Path) -> dict:
    reader = PdfReader(path)
    assert len(reader.pages) > 0
    metadata = reader.metadata or {}
    assert str(metadata.get("/Title", "")) == CONFIG["title"]
    assert str(metadata.get("/Author", "")) == CONFIG["author"]
    outlines = flatten_outline(reader.outline)
    assert any(CHAPTER_30 in title for title in outlines), "Bookmark Chapter 30 mancante"
    last_text = reader.pages[-1].extract_text() or ""
    assert FINAL_TEXT in last_text, "Frase finale assente dall'ultima pagina PDF"
    return {"pages": len(reader.pages), "bookmarks": len(outlines), "bytes": path.stat().st_size}


def find_opf(zf: zipfile.ZipFile) -> str:
    container = ET.fromstring(zf.read("META-INF/container.xml"))
    rootfile = container.find(".//{*}rootfile")
    assert rootfile is not None
    return rootfile.attrib["full-path"]


def inspect_epub(path: Path) -> dict:
    with zipfile.ZipFile(path) as zf:
        names = zf.namelist()
        assert names[0] == "mimetype"
        assert zf.getinfo("mimetype").compress_type == zipfile.ZIP_STORED
        assert zf.read("mimetype") == b"application/epub+zip"
        assert "META-INF/container.xml" in names
        opf_path = find_opf(zf)
        opf = ET.fromstring(zf.read(opf_path))
        metadata_text = " ".join((el.text or "") for el in opf.findall(".//{*}metadata/*"))
        assert CONFIG["title"] in metadata_text and CONFIG["author"] in metadata_text
        manifest = {item.attrib.get("id"): item.attrib.get("href") for item in opf.findall(".//{*}manifest/{*}item")}
        spine_ids = [item.attrib.get("idref") for item in opf.findall(".//{*}spine/{*}itemref")]
        readable = [manifest[item_id] for item_id in spine_ids if item_id in manifest and manifest[item_id].endswith((".xhtml", ".html"))]
        assert readable
        opf_dir = str(Path(opf_path).parent)
        last_path = str(Path(opf_dir) / readable[-1]) if opf_dir != "." else readable[-1]
        last_xhtml = zf.read(last_path).decode("utf-8")
        assert CHAPTER_30 in last_xhtml
        assert FINAL_TEXT in re.sub(r"<[^>]+>", " ", last_xhtml)
        xhtml_count = sum(1 for name in names if name.endswith((".xhtml", ".html")))
        assert any(name.endswith(".css") for name in names)
        return {"xhtml_documents": xhtml_count, "bytes": path.stat().st_size}


def main() -> int:
    paths = {ext: BUILD_DIR / f"{BASENAME}.{ext}" for ext in ("md", "docx", "pdf", "epub")}
    missing = [str(p.relative_to(ROOT)) for p in paths.values() if not p.exists()]
    if missing:
        print("Artifact mancanti: " + ", ".join(missing), file=sys.stderr)
        return 1
    results = {
        "Markdown": inspect_markdown(paths["md"]),
        "DOCX": inspect_docx(paths["docx"]),
        "PDF": inspect_pdf(paths["pdf"]),
        "EPUB": inspect_epub(paths["epub"]),
    }
    print("Build inspection PASS")
    for name, data in results.items():
        print(f"- {name}: " + ", ".join(f"{k}={v}" for k, v in data.items()))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
